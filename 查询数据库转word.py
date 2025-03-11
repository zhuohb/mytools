import os
import sys

import pymysql
from docx import Document

if sys.platform=='win32':
    os.system('chcp 65001')

# 连接 MySQL 数据库
connection = pymysql.connect(
    host='',
    user='',
    password='',
    db='',
    charset='utf8mb4',
    cursorclass=pymysql.cursors.DictCursor
)

try:
    with connection.cursor() as cursor:
        # 查询字段信息
        sql = '''
        SELECT 
            COLUMN_NAME AS '字段名',
            DATA_TYPE AS '类型',
            IS_NULLABLE AS '是否为空',
            COLUMN_DEFAULT AS '默认值',
            COLUMN_COMMENT AS '注释'
        FROM 
            INFORMATION_SCHEMA.COLUMNS
        WHERE 
            TABLE_SCHEMA = %s
            AND TABLE_NAME = %s
        '''
        cursor.execute(sql, ('sn21_33_017', 'tb_anjgl_cwcl'))
        result = cursor.fetchall()

        # 创建 Word 文档
        doc = Document()
        doc.add_heading('表字段信息', 0)

        # 添加表格
        table = doc.add_table(rows=1, cols=5)

        table.style = 'Table Grid'
        headers = ['字段名', '类型', '是否为空', '默认值', '注释']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        # 填充数据
        for row in result:
            new_row = table.add_row()
            new_row.cells[0].text = row['字段名']
            new_row.cells[1].text = row['类型']
            new_row.cells[2].text = row['是否为空']
            new_row.cells[3].text = str(row['默认值']) if row['默认值'] is not None else ''
            new_row.cells[4].text = row['注释']

        # 保存文档
        doc.save('table_info.docx')
        print("Word 文档已生成！")

finally:
    connection.close()
